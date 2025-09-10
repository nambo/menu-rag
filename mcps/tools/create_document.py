"""
创建word文档的方法

1. 图片自动插入到段落前
2. 生成下标去重与样式处理
3. 标题、段落格式处理

1.0 @nambo 2025-07-20
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from datetime import datetime
import json
import os
import re
import json
from PIL import Image

def combine_images_horizontally(paths, output_path=None):
    # 打开所有图片并转换为RGBA模式（兼容透明背景）
    images = [Image.open(path).convert("RGBA") for path in paths]
    
    # 计算拼接后的最大高度
    max_height = max(img.height for img in images)
    
    # 调整每张图片到相同高度（保持比例）
    resized_images = []
    for img in images:
        # 计算缩放后的新宽度（保持原始宽高比）
        w_percent = max_height / float(img.height)
        new_width = int(float(img.width) * w_percent)
        # 高质量缩放
        resized_img = img.resize((new_width, max_height), Image.LANCZOS)
        resized_images.append(resized_img)
    
    # 计算总宽度
    total_width = sum(img.width for img in resized_images)
    
    # 创建新画布（带透明背景）
    combined = Image.new("RGBA", (total_width, max_height), (0, 0, 0, 0))
    
    # 拼接图片
    x_offset = 0
    for img in resized_images:
        combined.paste(img, (x_offset, 0), img)
        x_offset += img.width
    
    # 保存结果（根据需求选择格式）
    if output_path is None or output_path == '':
        output_path = paths[0] + '.jpg'

    if output_path.lower().endswith('.jpg'):
        combined = combined.convert("RGB")  # JPG不支持透明，转为RGB

    combined.save(output_path)

    return output_path

def transform_data(a):
    all_source = []
    exist_source_map = {}
    all_source_map = {}
    
    # 第一部分：处理所有来源数据
    for a_i in a:
        sources = a_i.get('sources', [])
        for item in sources:
            # 生成唯一标识符
            item_id = f"{item.get('key', '')}{item.get('handler', '')}"
            old_idx = str(item.get('idx', ''))
            
            if item_id in exist_source_map:
                # 已存在时记录映射关系
                exist_item = exist_source_map[item_id]
                print(f"{old_idx} {exist_item.get('idx', '')}")
                all_source_map[old_idx] = exist_item['idx']
            else:
                # 新来源时分配新索引
                new_idx = str(len(all_source) + 1)
                all_source_map[str(item['idx'])] = new_idx
                item['idx'] = new_idx  # 修改原始对象
                item['summary'] = ''
                exist_source_map[item_id] = item
                all_source.append(item)
    
    # 第二部分：构建新的内容结构
    b = []
    for a_i in a:
        content_list = a_i.get('content_list', [])
        contents = []
        
        for cont in content_list:
            # 处理图片
            imgs_list = cont.get('imgs', [])
            r_imgs = [{'title': img.get('title'), 'path': img.get('path')} 
                      for img in imgs_list] if imgs_list else []
            
            # 处理文本内容
            txt = cont.get('content', '')
            for old_idx, new_idx in all_source_map.items():
                # 替换特定标签格式 [:<old_idx>]
                if f'[:{old_idx}]' in txt:
                    pattern = r'\[:' + re.escape(old_idx) + r'\]'
                    txt = re.sub(pattern, f'[:{new_idx}]', txt)
                    print(old_idx, '->', new_idx, txt)
            
            # 删除空标签 [::]
            txt = re.sub(r'\[\:\]', '', txt)

            if '\n' in txt:
                txt_list = txt.split('\n')
            else:
                txt_list = [txt]

            for idx, txt_item in enumerate(txt_list):
                if idx == 0:
                    contents.append({'txt': txt_item, 'imgs': r_imgs})
                else:
                    contents.append({'txt': txt_item, 'imgs': []})
        
        b.append({
            'title': a_i.get('title'),
            'contents': contents
        })
    
    # 构建最终结果
    res = {
        'content': b,
        'source': all_source
    }
    
    return res

def create_doc_with_references(content, doc):
    # 删除不带：的无效引用
    content = re.sub(r'\[\d+\]', '', content)
    
    # 匹配所有 [:数字] 的引用标记
    pattern = re.compile(r'\[:(\d+)\]')
    # 按引用标记分割内容
    parts = pattern.split(content)

    print('\n\n\n', parts, '\n\n\n')
    
    # 添加段落
    para = doc.add_paragraph()
    # 设置两端对齐
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # 设置首行缩进 2 字符（1 字符 ≈ 0.37 cm，2 字符 ≈ 0.74 cm）
    para.paragraph_format.first_line_indent = Cm(0.74)  # 或使用 Cm(0.74)
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # 添加普通文本部分
            if part:
                part = part.replace('**', '')
                if '#####' in part or '####' in part or '###' in part or '##' in part:
                    part = part.replace('#####', '').replace('####', '').replace('###', '').replace('##', '')
                    part = part.strip()
                    r = para.add_run(part)
                    if len(part) < 40:
                        r.font.bold = True
                else:
                    part = part.strip()
                    para.add_run(part)
        else:
            # 创建上标引用 [数字]
            run = para.add_run(f"[{part}]")
            run.font.superscript = True  # 设置为上标
            
            # 设置中文上标样式（可选）
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(9)  # 上标小一号字体

    return doc

def add_gray_border_to_run(run):
    """为Run中的图片添加灰色边框"""
    # 获取Run的XML元素
    r = run._r
    
    # 查找图片元素
    drawing = r.find('.//wp:inline', namespaces={'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
    if drawing is None:
        return  # 没有找到图片
    
    # 创建边框属性
    pic_pr = drawing.find('.//pic:picPr', namespaces={'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'})
    if pic_pr is None:
        # 如果picPr不存在，创建一个
        pic_pr = OxmlElement('pic:picPr')
        drawing.append(pic_pr)
    
    # 添加边框设置
    solid_fill = OxmlElement('a:solidFill')
    srgb_clr = OxmlElement('a:srgbClr')
    srgb_clr.set('val', '909090')  # 灰色
    solid_fill.append(srgb_clr)
    
    ln = OxmlElement('a:ln')
    ln.set('w', '7620')  # 边框宽度 7620 EMU ≈ 1.5pt
    ln.append(solid_fill)
    
    # 添加到图片属性
    pic_pr.append(ln)

def add_images_to_doc(doc, imgs):
    """添加图片到文档，满足布局要求"""
    # 单张图片处理
    if len(imgs) == 1:
        img_path = imgs[0]["path"]
        title = imgs[0]["title"]
        
        # 添加图片的段落
        p = doc.add_paragraph()
        
        # 添加图片
        try:
            run = p.add_run()
            shape = run.add_picture(img_path, width=Inches(5))
        except:
            return None
        
        p.paragraph_format.space_after = Pt(12)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        # 添加图片边框
        add_gray_border_to_run(run)
        
        # 添加标题
        title_p = doc.add_paragraph()
        title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_p.add_run(f"图表 {title}")
    
    # 两张图片处理
    elif len(imgs) == 2:
        # 创建单行两列的表格
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table.autofit = False
        
        # 设置表格宽度为页面宽度
        table_width = Inches(6.5)  # 根据页面调整
        table.columns[0].width = Inches(3)
        table.columns[1].width = Inches(3)
        
        # 添加图片到单元格
        for i, img in enumerate(imgs):
            cell = table.cell(0, i)
            cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 清除默认段落
            if cell.paragraphs:
                cell.paragraphs[0].clear()
            
            # 添加图片
            p = cell.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run()
            shape = run.add_picture(img["path"], height=Inches(2))  # 固定高度
            
            # 添加图片边框
            add_gray_border_to_run(run)
            
            # 添加标题
            title_p = cell.add_paragraph()
            title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_p.add_run(f"图表 {img['title']}")
        
        # 设置列间距（0.5cm）
        tblPr = table._tbl.tblPr
        spacing = OxmlElement('w:tblCellSpacing')
        spacing.set(qn('w:w'), str(int(Cm(0.5) * 567)))  # 转换为EMU
        spacing.set(qn('w:type'), 'dxa')
        tblPr.append(spacing)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

def generate_report(data, output_path, title_txt, data_date=None, target='', target_name=''):

    now_date = str(datetime.now().strftime("%Y-%m-%d-"))

    if data_date is None:
        data_date = str(datetime.now().strftime("%Y年%m月%d日"))
    else:
        data_date = datetime.strptime(data_date, "%Y-%m-%d").strftime("%Y年%m月%日")

    # 解析JSON数据
    data = transform_data(data)

    content = data["content"]
    sources = data["source"]
    
    # 创建Word文档
    doc = Document()
    
    # 设置全局样式
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    
    doc.add_paragraph("\n" * 3)  # 垂直居中留白

    # 1. 添加封面页
    # title = doc.add_heading('宁德时代（300750.SZ, 300750.HK）投资研究报告', 0)
    # title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # title_run = title.runs[0]
    # title_run.font.size = Pt(16)
    # title_run.font.bold = True
    # title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    cover = doc.add_paragraph()
    cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title = cover.add_run(title_txt)
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph("\n" * 18)  # 垂直居中留白
    
    institute = doc.add_paragraph()
    institute.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    institute_run = institute.add_run("发布机构：某某证券研究所")
    institute_run.font.size = Pt(12)
    
    date = doc.add_paragraph()
    date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date.add_run("分析师：某某某（执业证书编号：A123456789）​").font.size = Pt(12)

    date = doc.add_paragraph()
    date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date.add_run("报告时间：" + data_date).font.size = Pt(12)
    
    doc.add_page_break()
    
    # 2. 添加声明页
    disclaimer_title = doc.add_paragraph()
    disclaimer_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    disclaimer_title_run = disclaimer_title.add_run("重要声明与风险提示")
    disclaimer_title_run.font.size = Pt(16)
    disclaimer_title_run.font.bold = True
    
    doc.add_paragraph("\n" * 3)  # 垂直居中留白

    disclaimer = [
        f'1. 本公司未持有{target_name}相关证券股份；',
        '2. 本报告结论基于公开信息，不保证投资收益。投资者请谨慎决策，自主承担风险；',
        '3. 报告中的观点仅为分析师的独立判断，与本机构立场无关；',
        '4. 本报告版权归某某证券研究所所有，未经授权禁止转载。'
    ]
    
    for item in disclaimer:
        # p = doc.add_paragraph(style='List Bullet')
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.add_run(item)
    
    doc.add_paragraph("\n" * 10)

    date = doc.add_paragraph()
    date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date.add_run("报告撰写：某某某")

    date = doc.add_paragraph()
    date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date.add_run("合规审核：某某某（合规编号：H123456789）")

    date = doc.add_paragraph()
    date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date.add_run(" 存档编号：GFXXXXXXXX-001")
    
    doc.add_page_break()
    TITLE_IDX_MAP = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十', '十一', '十二', '十三', '十四', '十五']
    
    # 3. 添加目录
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    toc_title_run = toc_title.add_run("目    录")
    toc_title_run.font.size = Pt(16)
    toc_title_run.font.bold = True
    
    doc.add_paragraph("\n" * 3)  # 垂直居中留白
    # 自动生成目录（需实际文档生成后更新域）
    for idx, chapter in enumerate(content):
        # 添加章节标题
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = 2 * 10  # 层级缩进

        menu_title = TITLE_IDX_MAP[idx] + '、' + chapter["title"] + ' '
        menu_line = '-' * int((39 - len(menu_title)) / 0.55)
        
        runner = p.add_run(menu_title + menu_line + ' ' + str(idx*3 + 1))
        # runner.bold = True  # 一级标题加粗
        
        # 添加页码占位符
        # p.add_run("\t" + "...").underline = True  # 下划线模拟页码占位
    
    doc.add_page_break()
    
    # 4. 添加正文内容
    for idx, chapter in enumerate(content):
        # 添加章节标题
        heading = doc.add_heading(TITLE_IDX_MAP[idx] + '、' + chapter["title"], level=1)
        heading.runs[0].font.bold = True
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(12)

        for content_block in chapter["contents"]:
            txt = content_block["txt"]
            imgs = content_block["imgs"]
            
            # 添加图片
            # add_images_to_doc(doc, imgs)
            pre_img_title = ''
            for img in imgs:
                img_path = img["path"]

                # 如果是数组，则合并
                if isinstance(img_path, list):
                    img_path = combine_images_horizontally(img_path)

                try:
                    if not os.path.exists(img_path):
                        continue
                    img_title = img['title']

                    # 去重空的数据来源
                    if '（数据来源）' in img_title:
                        img_title = img_title.replace('（数据来源）', '')

                    # 去除连续两张相同的图片
                    if img_title == pre_img_title:
                        continue
                    pre_img_title = img_title

                    p = doc.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(5))
                    p.paragraph_format.space_after = Pt(12)

                    p = doc.add_paragraph()
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p.add_run(f"图表 {img_title}")
                except Exception as e:
                    print(f"图片插入失败: {e}")
            
            # 添加文本内容
            create_doc_with_references(txt, doc)
            # p = doc.add_paragraph(txt)
    
    doc.add_page_break()
    
    # 5. 添加参考文献
    ref_title = doc.add_heading("参考文献", level=1)
    ref_title.runs[0].font.bold = True
    ref_title.paragraph_format.space_before = Pt(12)
    ref_title.paragraph_format.space_after = Pt(12)
    
    for source in sources:
        source_date = source['date'] if 'date' in source else data_date.replace('年', '-').replace('月', '-').replace('日', '')
        source_name = source['source'].replace('\n', '').replace(' ', '').replace('<em>', '').replace('</em>', '')
        source_title = source['title'].replace('\n', '').replace(' ', '')
        source_txt = f"[{source['idx']}] {source_name} . {source_title} [EB/OL]. ({source_date})[{now_date}]. {source['url']}"
        doc.add_paragraph(
            source_txt
            # ,style='List Number'
        )
    
    # 6. 添加分析师声明
    # analyst_title = doc.add_heading("分析师声明", level=1)
    
    # analyst_declaration = [
    #     "本人具有中国证券业协会授予的证券投资咨询执业资格，本人承诺以专业审慎的态度、诚实守信的职业操守独立、客观地出具本报告。",
    #     "本报告清晰准确地反映了本人的研究观点，结论不受任何第三方的授意或影响。",
    #     "本人不曾因也将不会因本报告中的具体推荐意见或观点而直接或间接获取任何形式的补偿。"
    # ]
    
    # for statement in analyst_declaration:
    #     doc.add_paragraph(statement)
        
    # doc.add_paragraph("\n分析师：XXX")
    # doc.add_paragraph("签署日期：2025年7月20日")
    
    # 保存文档
    doc.save(output_path)
    print(f"研报已成功生成至: {output_path}")

    return output_path

# 使用示例
if __name__ == "__main__":
    with open('input_data.json', 'r', encoding='utf-8') as f:
        content_json = f.read()
        content_json = json.loads(content_json)
    
    output_path = "宁德时代投资分析报告1py.docx"
    generate_report(content_json, output_path, '宁德时代投资分析报告1py')